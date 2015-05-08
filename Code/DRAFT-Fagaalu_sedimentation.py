# -*- coding: utf-8 -*-
"""
Created on Wed Mar 18 06:26:25 2015

@author: Alex
"""
plt.ioff()
plt.close('all')
from docx import *
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
##  Create Document
document = Document()

######## SOME TOOLS
def add_figure_caption(fig_num=str(len(document.inline_shapes)),caption=''):
    cap = document.add_paragraph("Figure "+fig_num+". "+caption)
    cap.paragraph_style = 'caption'
    return
    
def dataframe_to_table(df=pd.DataFrame(),table_num=str(len(document.tables)+1),caption='',fontsize=11):
    table = document.add_table(rows=1, cols=len(df.columns)) 
    ## Merge all cells in top row and add caption text
    table_caption = table.rows[0].cells[0].merge(table.rows[0].cells[len(df.columns)-1])
    table_caption.text = "Table "+table_num+". "+caption
    ## Add  header
    header_row = table.add_row().cells
    col_count =0 ## counter  to iterate over columns
    for col in  header_row:
        col.text = df.columns[col_count] #df.columns[0] is the index
        col_count+=1
    ## Add data by  iterating over the DataFrame rows, then using a dictionary of DataFrame column labels to extract data
    col_labels = dict(zip(range(len(df.columns)),df.columns.tolist())) ## create dictionary where '1  to  n' is key for DataFrame columns
    for row in df.iterrows():  ## iterate over  rows in  DataFrame
        #print row[1]
        row_cells = table.add_row().cells ## Add a row to the  table
        col_count  = 0
        for cell in row_cells: ## iterate over the columns in the row
            #print row[1][str(col_labels[col_count])]
            cell.text = str(row[1][str(col_labels[col_count])]) ## and plug in data using a dictionary to  get column labels for DataFrame
            col_count+=1
    ## Format Table Style  
    table.style = 'TableGrid' 
    table.style.font.size = Pt(fontsize)
    table.autofit
    #table.num = str(len(document.tables)+1)
    return table
    
def add_equation(eq_table):
    t = document.add_table(rows=len(eq_table.rows),cols=len(eq_table.columns))
    t.rows[0].cells[1].text = eq_table.rows[0].cells[1].text
    t.rows[0].cells[2].text = eq_table.rows[0].cells[2].text   
    if len(eq_table.rows)>1:
        t.rows[1].cells[0].merge(t.rows[1].cells[2]).text = eq_table.rows[1].cells[0].text
    t.style = 'TableGrid'   
    t.autofit
    return t
###################################################################################################################################################################    


#### TABLES ########################################################################################################################################################
### Landcover_Table
table_count=0
def tab_count():
    global table_count
    table_count+=1
    return str(table_count)
# Prepare LULC Data
table_one = table_one_function()
table_one.table_num = str(tab_count())    


#### FIGURES ########################################################################################################################################################
figure_count=0
def fig_count():
    global figure_count
    figure_count+=1
    return str(figure_count)
## INTRODUCTION
#### Study Area Map
Study_Area_map = {'filename':maindir+'Figures/Maps/Study_Area.tif', 'fig_num':str(fig_count())}


###### EQUATIONS ############################################################################################################################################
equation_count=0
def eq_count():
    global equation_count
    equation_count+=1
    return str(equation_count)
    
Equations = Document(maindir+'/Manuscript/Equations.docx').tables

## SSYev = Q*SSC 
Equation_one = Equations[0].table
Equation_one.eq_num = eq_count()

##### APPENDIX #################################################################################################################################
#### Appendix
table_count,figure_count,equation_count=0, 0, 0
### Appendix Table One
Appendix_table_one = table_function() ## function to create table data
Appendix_table_one.table_num =str(tab_count())

############################################################################################################################################

#### TITLE
title_title = document.add_heading('TITLE:',level=1)
title_title.paragraph_format.space_before = 0
title = document.add_heading('Watershed and oceanic controls on spatial and temporal patterns of sediment deposition in a fringing reef embayment',level=1)
title.paragraph_format.space_before = 0
## subscript/superscript words
document.add_paragraph("")

#### ABSTRACT
abstract_title = document.add_heading('ABSTRACT',level=2)
abstract_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
abstract = document.add_paragraph('Abstract text goes here....')


#### INTRODUCTION
introduction_title = document.add_heading('Introduction',level=2)
document.add_paragraph("The complex spatial and temporal interaction of terrigenous sediment inputs, sediment resuspension, and hydrodynamic circulation can significantly alter the quantity, composition, and residence time of sediment in coral reefs which are significant controls on coral ecology (Storlazzi et al., 2009). Several studies correlate long term sediment accumulation, and resulting decreased coral health, with increased sediment supply from the watershed (Ryan et al., 2008, PNG studies), but there is also strong evidence of hydrodynamics decreasing sediment residence time in two ways: 1) by flushing suspended sediment away from the corals before it can be deposited (residence time = 0 min), and 2) resuspending and removing sediment that has been previously deposited. In contrast to many small, mountainous watersheds in temperate coastal regions where fluvial discharge and wave energy commonly coincide (Warrick et al., 2004), discharge, deposition, and reworking of flood sediment are often decoupled on tropical islands, causing high deposition rates and residence times of terrigenous sediment (Draut et al., 2009; Storlazzi et al., 2009). Conversely, seasonal wind and wave patterns in the tradewind belt can be coupled with sediment discharge or resuspension to decrease sediment deposition and residence times (Hoitink and Hoekstra, 2003; Muzuka et al., 2010). Given the increase in sediment discharge to coastal waters caused by anthropogenic watershed disturbance on tropical islands, an integrated understanding of how flood-supplied terrigenous sediment and water circulation control sediment deposition and residence time is essential for identifying and mitigating coral health impacts (Draut et al., 2009).")

document.add_paragraph("Many researchers and environmental managers are interested in determining the location and severity of terrigenous sediment impacts on coral health, but developing a measure of sediment impact has proven difficult. Several methods of measuring short-term (hourly to seasonal time scale) sediment accumulation have been developed, including discontinuous and quasi-continuous methods (Thomas 2004). Discontinuous methods require an observer to collect a sample on-site. Field collection is labor-intensive and only yields an average rate of accumulation over the observation interval (days to months), but allows the accumulated sediment to be further analyzed for composition, grain size, and pollutants. Quasi-continuous methods are more expensive and elaborate due to the use of electronic sensors and data loggers, but may be less labor intensive and yield sampling frequency only limited by the sensor and data logger, typically recording at frequencies of minutes for periods of several months. See Thomas and Ridd (2004) for a full review of measurement methods. Due to their relative cost and simplicity, tube traps are the most common method for measuring sediment accumulation in shallow coral reef environments (Storlazzi et al., 2011; White, 1990), but it is difficult to determine if these are ecologically meaningful indicators of coral stress. Some corals are well-adapted to turbid conditions (Perry et al., 2012), and deposited sediment can be removed actively by the coral itself (Barnes 1999), or passively by wave action before it is lethal. The stress on the coral organism increases linearly with the deposition amount and the duration of exposure (Fabricius, 2005) but tube traps overestimate deposition and do not allow for sediment resuspension, making it impossible to evaluate the residence time of deposited sediment (Storlazzi et al., 2011). Sediment can accumulate due to primary and secondary deposition of particlse (see Thomas 2004 for a review of terms), so to more accurately quantify “net” sedimentation, Field et al. (2012) proposed the use of “SedPods” where a flat surface allows for resuspension, similar to the surrounding benthic substrate. While the complex interaction of sediment composition, hydrodynamics, and coral physiology are important for the overall impacts on coral health, basic questions about location and controls on net terrigenous sediment accumulation rates are unknown at the study site.")

document.add_paragraph("Though some have shown good correlation between suspended sediment fluxes and rainfall (Meng 2008), several studies have found weak or no correlation between sediment trap collection and rainfall (Bothner et al., 2006; Victor et al., 2006) but it is well-known that SSY from small, mountainous watershed can be poorly correlated with precipitation (Basher et al., 2011; Duvert et al., 2012). Poor correlation between river discharge and sediment accumulation has also been found even in tidal estuaries (Pasternack and Brush 1998). By correlating sediment trap accumulation with measured and modeled SSY from the watershed, this research proposes to develop a model of spatially distributed, monthly sediment accumulation as a function of watershed inputs and hydrodynamic conditions. The proposed modeling approach is similar to other efforts that have attempted to limit the complexity of the modeling approach, but still account for the impact of ocean conditions on sediment dynamics (Fabricius et al. 2012).")

document.add_paragraph("The research questions for this paper are: How do flood-supplied terrigenous sediment and hydrodynamic conditions interact to control the gross and net rate of terrigenous sediment deposition at monthly time scales in a coral reef embayment? What controls the spatial distribution of sediment accumulation, and can it be predicted by the flow velocities of water over the reef and distance from the stream mouth?")

document.add_paragraph("In February and March, 2012, total sediment accumulation was measured at nine locations on Faga’alu reef using simple tube traps (TUBE), a ceramic tile (TILE), and an Astroturf mat (MAT). From April 2013 through June 2013 sediment accumulation was measured using SedPods. Sediment accumulation included both reef-derived carbonate and terrigenous sediment, and varied according to sediment trap type, location, and ocean conditions (Figure 3). The pilot study demonstrated strong spatial variability in sediment accumulation, with high rates near the stream mouth and on the northern reef. However, the pilot study data are insufficient to test the hypotheses due to a limited number of samples (n=5), no assessment of sediment composition, and insufficient data on hydrodynamic conditions. The methods developed in the pilot study have informed the development of the methods for the proposal.")

document.add_paragraph("")

#### STUDY AREA
study_area_title = document.add_heading('Study Area',level=2)
## Study Area map
if 'Study_Area_map' in locals():
    document.add_picture(Study_Area_map['filename'],width=Inches(6))
    add_figure_caption(Study_Area_map['fig_num'],"figure caption.")
    
#### METHODS
methods_title = document.add_heading('Methods',level=2)
document.add_heading('Measuring sediment accumulation on the reef',level=3)

document.add_paragraph("Deploying a TUBE in conjunction with a SedPod will allow comparison of gross and net sediment accumulation, and an assessment of the interaction of sediment loading and removal at time scales relevant to coral mortality and management. SedPods and TUBEs, deployed at nine locations on the reef flat (water depth 1-2 m) and reef crest (10-15 m) in Faga’alu Bay (Figure 3), are being collected monthly to provide data on sediment accumulation rates (g/m2/d) and composition from February 2014 through January 2015. Collection will be performed by Messina when in the field and by the Department of Marine and Wildlife Resources (DMWR) staff when Messina is not on-island. Sediment samples collected in tubes and SedPods will be wet sieved to the rinse salt from the sample and assess particle size (sand or fines). The samples will be dried and weighed to determine bulk sediment weight before being shipped to SDSU to characterize the geochemical composition (percent terrigenous, carbonate and organic) using Loss on Ignition (LOI) method (Heiri et al., 2001; Santisteban et al., 2004).")

document.add_heading('Modeling sediment accumulation',level=3)

document.add_paragraph("Statistical models, including both simple linear regression models and more complex generalized additive mixed models (GAMMs) will be used to establish the relative controls of each measured variable on sediment accumulation rates, both the average for North and South reefs, and at each of the nine locations where accumulation is measured. Sediment accumulation at location i (Si) during month t will be calculated:")

## Eq 1
document.add_paragraph("Sediment loading from the watershed in month t (Sw(t)) will be calculated using the model from Paper One:")
## Eq 2
document.add_paragraph("Water residence time for each 100m x 100m grid cell containing a TUBE/SedPod will be calculated from NOAA WW3 model output and the model developed in Paper Two. The relationship between swell height and residence time in each grid cell will be determined in Paper Two, of the form:")
## Eq 3
document.add_paragraph("Monthly sediment accumulation may be a function of sediment loading and hydrodynamic processes interacting on daily time scales, where hydrodynamic conditions only on the day of sediment discharge and not the mean monthly condition, are important. If monthly sediment loading and monthly mean residence time do not adequately predict sediment accumulation in the sediment traps, it might be necessary to investigate sediment loading and water residence times on daily scales, and further refine the statistical analysis and equations. In that case, daily sediment loading and daily mean residence time will be used to assess daily deposition, which can be compared to the monthly sediment accumulation measurements.")

document.add_heading('Temporal distribution of sediment accumulation',level=4)
document.add_paragraph("Two time scales of analysis will be used: monthly and seasonal (dry and wet season). A monthly time interval was chosen to correspond with other studies found in the literature (Muzuka et al., 2010; Victor et al., 2006), to include enough storm events to collect enough sediment for analysis, and for logistical reasons due to the high spatial coverage of sites and limited field personnel and resources. Assessing differences between dry and wet season sediment dynamics is useful to determine if there are seasonal patterns or modes that may be relevant to long term sediment accumulation (Ryan et al., 2008) or coral conservation and restoration (Muzuka et al., 2010). It is hypothesized that net deposition predominantly characterizes the wet season, and a net sediment removal, or limited deposition, predominantly characterizes in the dry season (Figure ).")

document.add_heading('Spatial distribution of sediment accumulation',level=4)
document.add_paragraph("An important consideration for coral conservation is determining the spatial distribution of sediment impacts. To explain the relative spatial variation of sediment accumulation among sediment traps, and to determine if flow direction or distance from the stream is more important, all sediment accumulation measurements will be normalized by the maximum of the measured accumulation at the nine traps for a given month. Normalized values are then modeled as a function of flow velocity (towards/away the stream mouth) and distance from the stream mouth:")

## Eq 4


document.add_paragraph("")

document
#### RESULTS ####
results_title = document.add_heading('Results and Discussion',level=2)
## Field data collection
document.add_heading('Field Data Collection',level=3)

#### CONCLUSION
conclusion_title=document.add_heading('Conclusion',level=2)
conclusion_text = document.add_paragraph("In conclusion...")

#### Appendix
document.add_page_break()
document.add_heading('APPENDIX',level=2)
## Appendix stuff goes here...

## Save Document
document.save(maindir+'Manuscript/DRAFT-Watershed and oceanic controls on spatial and temporal patterns of sediment deposition in a fringing reef embayment.docx')

## Clean up any open figures
plt.close('all')

